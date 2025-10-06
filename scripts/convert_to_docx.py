#!/usr/bin/env python3

import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file, docx_file):
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Split content by lines to process formatting
    lines = md_content.split('\n')
    
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
            
        # Handle headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:].strip(), level=3)
        # Handle lists
        elif line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            p = doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
        # Handle horizontal rules
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Regular paragraph with inline formatting
            p = doc.add_paragraph()
            
            # Process inline formatting
            text = line
            
            # Handle bold and links
            parts = re.split(r'(\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))', text)
            
            for part in parts:
                if not part:
                    continue
                    
                # Bold text
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                # Links
                elif '[' in part and '](' in part:
                    link_match = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
                    if link_match:
                        link_text = link_match.group(1)
                        # Just add the text (Word doesn't support hyperlinks easily in python-docx)
                        run = p.add_run(link_text)
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                else:
                    p.add_run(part)
    
    # Save the document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    markdown_to_docx("README.md", "Matthew_L_Resume.docx")